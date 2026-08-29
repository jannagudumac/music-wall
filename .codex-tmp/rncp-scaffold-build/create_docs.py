from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]


def set_font(run, name: str, size: int, color: str, bold: bool = False) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def build_document(path: Path, title: str, subtitle: str) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading = document.styles["Heading 1"]
    heading.font.name = "Calibri"
    heading._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading.font.size = Pt(16)
    heading.font.color.rgb = RGBColor(46, 116, 181)
    heading.paragraph_format.space_before = Pt(16)
    heading.paragraph_format.space_after = Pt(8)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(72)

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_after = Pt(8)
    set_font(title_paragraph.add_run(title), "Calibri", 30, "123E48", True)

    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_paragraph.paragraph_format.space_after = Pt(18)
    set_font(subtitle_paragraph.add_run(subtitle), "Calibri", 18, "F56B61")

    status = document.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status.paragraph_format.space_after = Pt(72)
    set_font(status.add_run("Document de travail — à compléter"), "Calibri", 11, "667A80")

    document.add_heading("Structure à compléter", level=1)
    document.add_paragraph(
        "Cette trame est prête à recevoir le contenu final du dossier RNCP Music Wall."
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("Music Wall • RNCP"), "Calibri", 9, "667A80")

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


build_document(
    ROOT / "RNCP" / "01_Dossier_Projet" / "dossier-projet.docx",
    "Dossier de projet",
    "Music Wall",
)
build_document(
    ROOT / "RNCP" / "02_Dossier_Professionnel" / "dossier-professionnel.docx",
    "Dossier professionnel",
    "Parcours RNCP",
)
